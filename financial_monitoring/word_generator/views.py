import os
from datetime import date
from rest_framework.response import Response
from rest_framework.views import APIView
from docx import Document
from io import BytesIO
from django.http import HttpResponse
from general_info.serializers import GeneralInfoSerializer
from general_info.models import GeneralInfo
from photo.models import Photo
from personal_data.models import PersonalData
from datetime import datetime
from personal_data.serializers import PersonalDataSerializer
from education.serializers import EducationSerializer
from academic_degree.serializers import AcademicDegreeSerializer
from working_history.serializers import WorkingHistorySerializer
from orders_list.serializers import OrdersListSerializer
from photo.serializers import PhotoSerializer
from docx.shared import Pt, Cm
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENTATION
from docx.shared import Inches
import base64


class GenerateDocumentView(APIView):
    available_serializers = {
        "personal_data": PersonalDataSerializer,
        "educations": EducationSerializer,
        "academic_degree": AcademicDegreeSerializer,
        "working_histories": WorkingHistorySerializer,
        "photos": PhotoSerializer,
        "orders_list": OrdersListSerializer,
    }

    # month_names = {
    #     "1": "Января",
    #     "2": "Февраля",
    #     "3": "Марта",
    #     "4": "Апреля",
    #     "5": "Мая",
    #     "6": "Июня",
    #     "7": "Июля",
    #     "8": "Августа",
    #     "9": "Сентября",
    #     "10": "Октября",
    #     "11": "Ноября",
    #     "12": "Декабря",
    # }

    def post(self, request, id, *args, **kwargs):
        general_info = GeneralInfo.objects.get(id=id)

        template_path = "word_generator/static/templates/spravka_template.docx"
        template_doc = Document(template_path)

        general_info_serializer = GeneralInfoSerializer(general_info)

        for related_name, serializer in self.available_serializers.items():
            if related_name == "personal_data":
                personal_data = getattr(general_info, related_name).all()
            elif related_name == "educations":
                educations_queryset = getattr(general_info, related_name).all()
            elif related_name == "academic_degree":
                academic_degree_queryset = getattr(general_info, related_name).all()
            elif related_name == "working_histories":
                working_histories_queryset = getattr(general_info, related_name).all()
            elif related_name == "photos":
                photos_queryset = getattr(general_info, related_name).all()
            elif related_name == "orders_list":
                orders_list_queryset = getattr(general_info, related_name).all()

        personal_data_serialized = PersonalDataSerializer(personal_data, many=True).data
        personal_data_data = [dict(item) for item in personal_data_serialized]
        educations_serialized = EducationSerializer(educations_queryset, many=True).data
        educations_data = [dict(item) for item in educations_serialized]
        academic_degree_serialized = AcademicDegreeSerializer(academic_degree_queryset, many=True).data
        academic_degree_data = [dict(item) for item in academic_degree_serialized]
        working_histories_serialized = WorkingHistorySerializer(working_histories_queryset, many=True).data
        working_data = [dict(item) for item in working_histories_serialized]
        photos = PhotoSerializer(photos_queryset, many=True).data
        photos_data = [dict(item) for item in photos]
        #retrieving the first photo
        image_data = base64.b64decode(photos_data[0]['photo'])

        orders_list_serialized = OrdersListSerializer(orders_list_queryset, many=True).data
        orders_list_data = [dict(item) for item in orders_list_serialized]

        name = general_info_serializer.data["firstname"]
        surname = general_info_serializer.data["surname"]
        patronymic = general_info_serializer.data["patronymic"]
        birth_date = datetime.strptime(general_info_serializer.data["birth_date"], "%Y-%m-%d").date()
        birth_day = str(birth_date.day)
        if birth_date.day < 10:
            birth_day = f"0{birth_day}"
        birth_month = str(birth_date.month)
        if birth_date.month < 10:
            birth_month = f"0{birth_month}"

        birth_year = str(birth_date.year)
        iin = general_info_serializer.data["iin"]
        birth_oblast = general_info_serializer.data["birth_oblast"]
        birth_city = general_info_serializer.data["birth_city"]
        birth_region = ''
        region_optional = general_info_serializer.data["birth_region"]
        if region_optional is not None:
            birth_region = f", регион {region_optional}"

        nationality = general_info_serializer.data["nationality"]

        dolzhnost = "не имеет"
        if personal_data_serialized:
            dolzhnost = personal_data_serialized[0]["jposition"]

        educations = "не имеет"
        if educations_data:
            text = ""
            for education in educations_data:
                out_year = datetime.strptime(education["education_date_out"], "%Y-%m-%d").date().year
                text += f"{education['education_type']}, в {out_year} году окончил(а) {education['education_place']} ({education['education_speciality']}). "
            educations = text

        academic_degree = "не имеет"
        if academic_degree_data:
            text = ""
            for degree in academic_degree_data:
                out_year = datetime.strptime(degree["diploma_date"], "%Y-%m-%d").date().year
                text += f"в {out_year} получил степень {degree['academic_degree']} в {degree['education_place']}. "
            academic_degree = text

        name_placehoder = "name"
        surname_placehoder = "surname"
        patronymic_placeholder = "patronymic"
        birth_day_placeholder = "day"
        birth_month_placeholder = "month"
        birth_year_placeholder = "year"
        iin_placeholder = "iin"
        birth_oblast_placeholder = "oblast"
        birth_city_placeholder = "city"
        birth_region_placeholder = "region"
        nationality_placeholder = "nationality"
        dolzhnost_placeholder = "position"
        educations_placeholder = "education"
        academic_degree_placeholder = "academicdegree"
        image_placeholder = "photo"

        current_job_start_date = ''
        for orders in orders_list_data:
            if orders["order_type"] == "О назначении":
                order_date = datetime.strptime(orders['order_date'], "%Y-%m-%d")
                order_date_formatted = order_date.strftime("%d.%m.%Y")
                current_job_start_date = order_date_formatted

        table_data = []
        for education in educations_data:
            table_data.append({"date_in": education["education_date_in"], "value": education})

        for work in working_data:
            table_data.append({"date_in": work["working_start"], "value": work})

        def custom_sort_key(item):
            date_str = item.get("date_in")
            return datetime.strptime(date_str, "%Y-%m-%d")

        sorted_table_data = sorted(table_data, key=custom_sort_key)

        if len(sorted_table_data) != 0:
            table = template_doc.add_table(rows=len(sorted_table_data), cols=3)
            table.style = 'Table Grid'

            for i, row in enumerate(table.rows):
                if i < len(sorted_table_data):
                    period_start, period_end, details = row.cells

                    fields_dict = sorted_table_data[i]['value']

                    if "working_start" in fields_dict:
                        paragraph_period = period_start.paragraphs[0]
                        start_date = datetime.strptime(fields_dict['working_start'], "%Y-%m-%d")
                        start_date_formatted = start_date.strftime("%d.%m.%Y")
                        paragraph_period.add_run(str(start_date_formatted))
                        run = paragraph_period.runs[0]
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(12)
                        period_start.width = Pt(100)

                        paragraph_period = period_end.paragraphs[0]
                        end_date = datetime.strptime(fields_dict["working_end"], "%Y-%m-%d")
                        end_date_formatted = end_date.strftime("%d.%m.%Y")
                        paragraph_period.add_run(str(end_date_formatted))
                        run = paragraph_period.runs[0]
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(12)
                        period_end.width = Pt(100)

                        paragraph_details = details.paragraphs[0]
                        paragraph_details.add_run(
                            f"{fields_dict['jposition']} {fields_dict['orfanization_name']}")
                        run = paragraph_details.runs[0]
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(12)
                        details.width = Pt(1000)
                    else:
                        paragraph_period = period_start.paragraphs[0]
                        start_date = datetime.strptime(fields_dict['education_date_in'], "%Y-%m-%d")
                        start_date_formatted = start_date.strftime("%d.%m.%Y")
                        paragraph_period.add_run(str(start_date_formatted))
                        run = paragraph_period.runs[0]
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(12)
                        period_start.width = Pt(100)

                        paragraph_period = period_end.paragraphs[0]
                        end_date = datetime.strptime(fields_dict["education_date_out"], "%Y-%m-%d")
                        end_date_formatted = end_date.strftime("%d.%m.%Y")
                        paragraph_period.add_run(str(end_date_formatted))
                        run = paragraph_period.runs[0]
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(12)
                        period_end.width = Pt(100)

                        paragraph_details = details.paragraphs[0]
                        paragraph_details.add_run(f"Учился, {fields_dict['education_place']} ({fields_dict['education_speciality']})")
                        run = paragraph_details.runs[0]
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(12)
                        details.width = Pt(1000)
        else:
            table = template_doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'

        if len(sorted_table_data) != 0:
            period_start, period_end, details = table.add_row().cells

            paragraph_period = period_start.paragraphs[0]
            paragraph_period.add_run(current_job_start_date)
            run = paragraph_period.runs[0]
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            period_start.width = Pt(100)

            paragraph_period = period_end.paragraphs[0]
            paragraph_period.add_run("По н/вр")
            run = paragraph_period.runs[0]
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            period_end.width = Pt(100)

            paragraph_details = details.paragraphs[0]
            paragraph_details.add_run(
                f"{dolzhnost} {personal_data_data[0]['departament']}")
            run = paragraph_details.runs[0]
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            details.width = Pt(1000)



        for paragraph in template_doc.paragraphs:
            for run in paragraph.runs:
                if birth_region is not None:
                    run.text = run.text.replace(birth_region_placeholder, birth_region)
                else:
                    run.text = run.text.replace(birth_region_placeholder, '')

                run.text = run.text.replace(surname_placehoder, surname)
                run.text = run.text.replace(name_placehoder, name)
                run.text = run.text.replace(patronymic_placeholder, patronymic)
                run.text = run.text.replace(
                    birth_day_placeholder, birth_day)
                run.text = run.text.replace(
                    birth_month_placeholder, birth_month)
                run.text = run.text.replace(
                    birth_year_placeholder, birth_year)
                run.text = run.text.replace(
                    iin_placeholder, iin)
                run.text = run.text.replace(
                    birth_oblast_placeholder, birth_oblast)
                run.text = run.text.replace(
                    birth_city_placeholder, birth_city)
                run.text = run.text.replace(
                    nationality_placeholder, nationality)
                run.text = run.text.replace(
                    dolzhnost_placeholder, dolzhnost)
                run.text = run.text.replace(
                    educations_placeholder, educations)
                run.text = run.text.replace(
                    academic_degree_placeholder, academic_degree)

        modified_template_buffer = BytesIO()

        image_placeholder = template_doc.paragraphs[0].runs[0].add_picture(BytesIO(image_data))

        image_placeholder.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        image_placeholder.space_before = Inches(2)

        image_placeholder.width = Cm(3)
        image_placeholder.height = Cm(4)

        template_doc.save(modified_template_buffer)
        modified_template_buffer.seek(0)

        response = HttpResponse(
            modified_template_buffer,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename=result.docx'

        return response

        # user = GeneralInfoSerializer(GeneralInfo.objects.get(id=id))
        #
        # template_path = "word_generator/static/templates/template.docx"
        # template_doc = Document(template_path)
        #
        # date1 = datetime.strptime(
        #     request.data["date_depart"], "%Y-%m-%d").date()
        # date2 = datetime.strptime(
        #     request.data["date_return"], "%Y-%m-%d").date()
        #
        # date_difference = date2 - date1
        # number_of_days = str(date_difference.days)
        #
        # name = user.data["firstname"]
        # surname = user.data["surname"]
        # patronymic = user.data["patronymic"]
        # order_type = request.data.get('order_type', None)
        # type_of_order_type = request.data.get('type_of_order_type', None)
        #
        # day_depart = str(date1.day)
        # month_depart = self.month_names[str(date1.month)]
        # year_depart = str(date1.year)
        #
        # day_return = str(date2.day)
        # month_return = self.month_names[str(date2.month)]
        # year_return = str(date2.year)
        #
        # current_date = datetime.now()
        # current_day = str(current_date.day)
        # current_month = self.month_names[str(current_date.month)]
        # current_year = str(current_date.year)
        #
        # name_placehoder = "name"
        # surname_placehoder = "surname"
        # patronymic_placeholder = "patronymic"
        # order_type_placehoder = "ordertype"
        # type_of_order_type_placeholder = "typeof"
        # counts_day_placeholder = "count"
        #
        # day_depart_placeholder = "daydepart"
        # month_depart_placeholder = "monthdepart"
        # year_depart_placeholder = "yeardepart"
        #
        # day_return_placeholder = "dayreturn"
        # month_return_placeholder = "monthreturn"
        # year_return_placeholder = "yearreturn"
        #
        # day_now_placeholder = "daynow"
        # month_now_placeholder = "monthnow"
        # year_now_placeholder = "yearnow"
        #
        # for paragraph in template_doc.paragraphs:
        #     for run in paragraph.runs:
        #         run.text = run.text.replace(surname_placehoder, surname)
        #         run.text = run.text.replace(name_placehoder, name)
        #         run.text = run.text.replace(patronymic_placeholder, patronymic)
        #         run.text = run.text.replace(order_type_placehoder, order_type)
        #
        #         if type_of_order_type is not None:
        #             run.text = run.text.replace(
        #                 type_of_order_type_placeholder, type_of_order_type)
        #         else:
        #             run.text = run.text.replace(
        #                 type_of_order_type_placeholder, "")
        #
        #         run.text = run.text.replace(
        #             counts_day_placeholder, number_of_days)
        #
        #         run.text = run.text.replace(day_depart_placeholder, day_depart)
        #         run.text = run.text.replace(
        #             month_depart_placeholder, month_depart)
        #         run.text = run.text.replace(
        #             year_depart_placeholder, year_depart)
        #
        #         run.text = run.text.replace(day_return_placeholder, day_return)
        #         run.text = run.text.replace(
        #             month_return_placeholder, month_return)
        #         run.text = run.text.replace(
        #             year_return_placeholder, year_return)
        #
        #         run.text = run.text.replace(day_now_placeholder, current_day)
        #         run.text = run.text.replace(
        #             month_now_placeholder, current_month)
        #         run.text = run.text.replace(year_now_placeholder, current_year)
        #
        # modified_template_buffer = BytesIO()
        # template_doc.save(modified_template_buffer)
        # modified_template_buffer.seek(0)
        #
        # response = HttpResponse(
        #     modified_template_buffer,
        #     content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        # )
        # response['Content-Disposition'] = 'attachment; filename=Modified_Template.docx'
        #
        # return response

    # def get(self, request, *args, **kwargs):
    #     template_path = "word_generator/static/templates/template.docx"

    #     with open(template_path, 'rb') as template_file:
    #         template_content = template_file.read()
    #         print(template_content)

    #     response = HttpResponse(template_content,
    #                             content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    #     response['Content-Disposition'] = 'attachment; filename=Template.docx'

    #     return response
